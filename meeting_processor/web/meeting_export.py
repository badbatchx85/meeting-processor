"""Export a single meeting's summary to Markdown or Word (.docx).

Source of truth is the meeting's Resumo note (already parsed by
``_load_meeting`` into ``{meta, resumo_md, tasks}``). The transcript is
intentionally excluded — it is large and viewable separately.
"""
from __future__ import annotations

import io
import re

_TRANSCRIPT_HEADING = "## Transcricao Completa"


def to_markdown(meeting: dict) -> str:
    """Return the summary body, trimming the trailing Obsidian transcript link."""
    body = meeting.get("resumo_md", "") or ""
    idx = body.find(_TRANSCRIPT_HEADING)
    if idx != -1:
        body = body[:idx]
    return body.rstrip() + "\n"


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, rendering **bold** segments as bold runs."""
    for i, segment in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = i % 2 == 1  # odd indices are the captured bold groups


def to_docx(meeting: dict) -> bytes:
    """Render the summary markdown (from ``to_markdown``) into a .docx byte stream."""
    from docx import Document

    md = to_markdown(meeting)
    doc = Document()

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- [ ] ") or line.startswith("- [x] "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[6:].strip())
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
        elif line.startswith(">"):
            continue  # skip Obsidian callouts/quotes
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
