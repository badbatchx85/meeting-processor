"""Tests for the meetings list new fields + markdown/docx export."""
import io


def _write_meeting(vault_path):
    """Create a meeting folder with a Resumo note carrying the new fields."""
    folder = vault_path / "wiki" / "reunioes" / "2026-06-04 10h00 - reuniao"
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "Resumo - 2026-06-04 10h00 - reuniao.md").write_text(
        """---
type: source
title: "2026-06-04 10h00 - reuniao"
created: 2026-06-04
duration: "10m"
participants: ["Ana"]
meeting_type: "planejamento"
purpose: "Alinhar o roadmap"
---

# 2026-06-04 10h00 - reuniao

**Tipo:** planejamento
**Participantes:** Ana

## Propósito

Alinhar o roadmap

## Resumo Executivo

Discussão do roadmap.

## Decisões

- Adiar o lançamento

## Tarefas Identificadas

- [ ] Preparar deck (Responsavel: Ana)

## Questões em Aberto

- Quem assume o suporte?

## Transcricao Completa

> [!info] Transcricao original
> [[Transcricao - 2026-06-04 10h00 - reuniao|Transcricao]]
""",
        encoding="utf-8",
    )
    return folder.name


def test_list_meetings_includes_type_and_purpose(client, config):
    _write_meeting(config.vault_path)
    r = client.get("/api/meetings")
    assert r.status_code == 200
    m = r.json()[0]
    assert m["meeting_type"] == "planejamento"
    assert m["purpose"] == "Alinhar o roadmap"


def test_export_md_returns_summary_without_transcript_link(client, config):
    mid = _write_meeting(config.vault_path)
    r = client.get(f"/api/meetings/{mid}/export.md")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/markdown")
    assert "attachment" in r.headers.get("content-disposition", "")
    body = r.text
    assert "Alinhar o roadmap" in body          # purpose
    assert "- Adiar o lançamento" in body        # decision
    assert "Quem assume o suporte?" in body       # open question
    assert "## Transcricao Completa" not in body  # transcript link stripped


def test_export_md_missing_meeting_404(client):
    r = client.get("/api/meetings/nope/export.md")
    assert r.status_code == 404


def test_export_docx_returns_valid_document(client, config):
    import docx  # python-docx

    mid = _write_meeting(config.vault_path)
    r = client.get(f"/api/meetings/{mid}/export.docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in r.headers.get("content-disposition", "")
    assert len(r.content) > 0

    doc = docx.Document(io.BytesIO(r.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Alinhar o roadmap" in full_text                 # purpose present
    assert any(p.style.name.startswith("Heading") for p in doc.paragraphs)  # a heading
    assert "Adiar o lançamento" in full_text                # decision bullet


def test_export_docx_missing_meeting_404(client):
    r = client.get("/api/meetings/nope/export.docx")
    assert r.status_code == 404


def test_export_roundtrip_from_real_note(client, config):
    from datetime import datetime
    from meeting_processor.models import (
        ActionItem, MeetingSummary, Transcript, TranscriptSegment,
    )
    from meeting_processor.note_generator import NoteGenerator

    transcript = Transcript(
        segments=[TranscriptSegment(start=0.0, end=5.0, text="Olá pessoal.")],
        full_text="Olá pessoal.", language="pt", duration=5.0,
    )
    summary = MeetingSummary(
        executive_summary="Resumo.",
        time_windows=[],
        action_items=[ActionItem(description="Preparar deck")],
        participants=["Ana"],
        key_topics=["Roadmap"],
        purpose="Alinhar o roadmap",
        meeting_type="planejamento",
        decisions=["Adiar o lançamento"],
        open_questions=["Quem assume o suporte?"],
    )
    gen = NoteGenerator(config)
    created = datetime(2026, 6, 4, 10, 0)
    paths = gen.prepare("reuniao.mp4", created)
    gen.write_transcription(transcript, paths)
    gen.write_summary_note(transcript, summary, "reuniao.mp4", created, paths)
    gen.write_group_note(paths, has_summary=True)

    mid = paths.folder_name

    r = client.get(f"/api/meetings/{mid}/export.md")
    assert r.status_code == 200
    body = r.text
    assert "Alinhar o roadmap" in body
    assert "## Decisões" in body
    assert "- Adiar o lançamento" in body
    assert "## Questões em Aberto" in body
    assert "## Transcricao Completa" not in body
    assert "{tarefas_link}" not in body  # f-string fix holds end-to-end

    r2 = client.get(f"/api/meetings/{mid}/export.docx")
    assert r2.status_code == 200
    assert r2.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
